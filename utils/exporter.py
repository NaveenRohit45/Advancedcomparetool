from docx import Document
from io import BytesIO

def export_to_docx(diffs, entities, tables, tones, summary):
    doc = Document()
    doc.add_heading("Document Comparison Report", 0)
    doc.add_heading("Text Differences", level=1)
    doc.add_paragraph(diffs, style="Normal")
    doc.add_heading("Entity Differences", level=1)
    doc.add_paragraph(str(entities))
    doc.add_heading("Table Differences", level=1)
    doc.add_paragraph(str(tables))
    doc.add_heading("Tone Analysis", level=1)
    doc.add_paragraph(str(tones))
    doc.add_heading("AI Summary", level=1)
    doc.add_paragraph(summary)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
